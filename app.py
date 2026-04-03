import os
import io
import csv
import random
import string
import zipfile
import tempfile
import shutil
from datetime import datetime, timedelta
from flask import Flask, render_template, request, redirect, url_for, session, send_file, jsonify, flash
from flask_sqlalchemy import SQLAlchemy
from flask_cors import CORS
from flask_migrate import Migrate
from flask_babel import Babel
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from docx import Document
from textblob import TextBlob
from sqlalchemy import text
from sqlalchemy.orm import joinedload

# --- Environment Configuration ------------------------------------
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'gweru_secret_key_2026')
# SQLite with timeout and thread-safe mode to prevent database lock errors
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///feedback.db?timeout=20&check_same_thread=False')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
    'pool_recycle': 300,
    'pool_size': 1,
    'max_overflow': 0
}
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB
app.config['BABEL_DEFAULT_LOCALE'] = 'en'
app.config['BABEL_TRANSLATION_DIRECTORIES'] = 'translations'

def create_upload_folder():
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

create_upload_folder()

# --- Extensions ---------------------------------------------------
db = SQLAlchemy(app)
migrate = Migrate(app, db)
CORS(app, resources={r"/*": {"origins": "*"}})
babel = Babel()
babel.init_app(app, locale_selector=lambda: session.get('lang', request.accept_languages.best_match(['en', 'sn', 'nd'])))

# --- Helper Functions ---------------------------------------------
def generate_reference():
    date_part = datetime.now().strftime('%Y%m%d')
    random_part = ''.join(random.choices(string.ascii_uppercase + string.digits, k=4))
    return f"FB-{date_part}-{random_part}"

def log_audit(action, details=None):
    user_id = session.get('user_id')
    log = AuditLog(user_id=user_id, action=action, details=details)
    db.session.add(log)
    db.session.commit()

def parse_datetime(dt_str):
    if not dt_str:
        return None
    try:
        return datetime.strptime(dt_str, '%Y-%m-%d %H:%M:%S.%f')
    except ValueError:
        return datetime.strptime(dt_str, '%Y-%m-%d %H:%M:%S')

# --- Admin decorator ----------------------------------------------
def admin_required(f):
    from functools import wraps
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not session.get('logged_in') or session.get('role') != 'admin':
            flash('Admin access required.', 'error')
            return redirect(url_for('staff_dashboard'))
        return f(*args, **kwargs)
    return decorated_function

# --- NLTK Data ----------------------------------------------------
import nltk
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)

# -----------------------------------------------------------------
# DATABASE MODELS
# -----------------------------------------------------------------
feedback_categories = db.Table('feedback_categories',
    db.Column('feedback_id', db.Integer, db.ForeignKey('feedback.id'), primary_key=True),
    db.Column('category_id', db.Integer, db.ForeignKey('category.id'), primary_key=True)
)

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(200), nullable=False)
    role = db.Column(db.String(50), nullable=False, default='viewer')

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Category(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), unique=True, nullable=False)

class Feedback(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    reference = db.Column(db.String(20), unique=True, nullable=False, default=generate_reference)
    issue_received = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.now)
    type = db.Column(db.String(50), nullable=False)
    mechanism = db.Column(db.String(100), nullable=True)
    recommendation = db.Column(db.String(500), nullable=True)
    first_action = db.Column(db.String(500), nullable=True)
    action_taken_at = db.Column(db.DateTime, nullable=True)
    implementation_status = db.Column(db.String(100), nullable=True)
    final_comment = db.Column(db.String(500), nullable=True)
    final_status = db.Column(db.String(100), nullable=True)
    action_timestamp = db.Column(db.DateTime, nullable=True)
    resolved_at = db.Column(db.DateTime, nullable=True)
    contact_email = db.Column(db.String(120), nullable=True)
    contact_phone = db.Column(db.String(20), nullable=True)

    categories = db.relationship('Category', secondary=feedback_categories, lazy='subquery',
                                 backref=db.backref('feedbacks', lazy=True))
    history = db.relationship('FeedbackHistory', backref='feedback', lazy=True, cascade='all, delete-orphan')
    attachments = db.relationship('Attachment', backref='feedback', lazy=True, cascade='all, delete-orphan')

    def time_taken_to_resolve(self):
        if self.resolved_at:
            return (self.resolved_at - self.created_at).days
        return None

    def time_to_action(self):
        if self.action_taken_at:
            return self.action_taken_at - self.created_at
        return None

    def sentiment(self):
        blob = TextBlob(self.issue_received)
        return blob.sentiment.polarity

class FeedbackHistory(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    feedback_id = db.Column(db.Integer, db.ForeignKey('feedback.id'), nullable=False)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    action = db.Column(db.String(50), nullable=False)
    old_status = db.Column(db.String(100), nullable=True)
    new_status = db.Column(db.String(100), nullable=True)
    comment = db.Column(db.String(500), nullable=True)
    timestamp = db.Column(db.DateTime, default=datetime.now)

class Attachment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    feedback_id = db.Column(db.Integer, db.ForeignKey('feedback.id'), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    filepath = db.Column(db.String(500), nullable=False)
    uploaded_at = db.Column(db.DateTime, default=datetime.now)

class AuditLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    action = db.Column(db.String(100), nullable=False)
    details = db.Column(db.Text, nullable=True)
    timestamp = db.Column(db.DateTime, default=datetime.now)

    user = db.relationship('User', backref='audit_logs')

# --- Database initialization (auto-create tables & admin) ---
with app.app_context():
    db.create_all()
    # Create admin user if none exists
    if not User.query.filter_by(username="admin").first():
        admin = User(
            username="admin",
            password_hash=generate_password_hash("admin123"),
            role="admin"
        )
        db.session.add(admin)
        db.session.commit()
        print("✅ Database initialized and admin user created (username: admin, password: admin123)")
    else:
        print("✅ Database already initialized.")

# -----------------------------------------------------------------
# AUTHENTICATION ROUTES
# -----------------------------------------------------------------
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]
        user = User.query.filter_by(username=username).first()
        if user and user.check_password(password):
            session["user_id"] = user.id
            session["username"] = user.username
            session["role"] = user.role
            session["logged_in"] = True
            log_audit('login', f'User {username} logged in')
            return redirect(url_for("staff_dashboard"))
        else:
            return render_template("login.html", error="Invalid credentials.")
    return render_template("login.html")

@app.route("/logout")
def logout():
    log_audit('logout', f'User {session.get("username")} logged out')
    session.clear()
    return redirect(url_for("home"))

@app.before_request
def require_login():
    public_endpoints = ["login", "static", "anonymous_feedback", "anonymous_thanks",
                        "public_dashboard", "track_feedback", "about", "contact",
                        "help", "home"]
    if request.endpoint not in public_endpoints and not session.get("logged_in"):
        return redirect(url_for("login"))

# -----------------------------------------------------------------
# PUBLIC ROUTES
# -----------------------------------------------------------------
@app.route("/")
def home():
    return render_template("welcome.html")

@app.route("/anonymous", methods=["GET", "POST"])
def anonymous_feedback():
    if request.method == "POST":
        uploaded_files = request.files.getlist("attachments")
        filenames = []
        for file in uploaded_files:
            if file and file.filename:
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                filenames.append(filename)

        fb = Feedback(
            issue_received=request.form["issue"],
            type=request.form["type"],
            mechanism="Anonymous Online",
            recommendation=request.form.get("recommendation"),
            contact_email=request.form.get("contact_email"),
            contact_phone=request.form.get("contact_phone")
        )
        db.session.add(fb)
        db.session.flush()

        for fname in filenames:
            att = Attachment(feedback_id=fb.id, filename=fname, filepath=fname)
            db.session.add(att)

        db.session.commit()
        log_audit('anonymous_feedback_submitted', f'Feedback {fb.reference} submitted')
        return render_template("anonymous_thanks.html", reference=fb.reference)
    return render_template("anonymous_feedback.html")

@app.route("/public")
def public_dashboard():
    total_feedback = Feedback.query.count()
    complaints = Feedback.query.filter_by(type='complaint').count()
    suggestions = Feedback.query.filter_by(type='suggestion').count()
    compliments = Feedback.query.filter_by(type='compliment').count()
    resolved_rate = 0
    if complaints > 0:
        resolved = Feedback.query.filter_by(type='complaint', final_status='Resolved').count()
        resolved_rate = round(resolved / complaints * 100, 1)
    monthly = db.session.query(
        db.func.strftime('%Y-%m', Feedback.created_at).label('month'),
        db.func.count().label('count')
    ).group_by('month').order_by('month').all()
    months = [m[0] for m in monthly]
    counts = [m[1] for m in monthly]
    return render_template("public_dashboard.html",
                           total=total_feedback,
                           complaints=complaints,
                           suggestions=suggestions,
                           compliments=compliments,
                           resolved_rate=resolved_rate,
                           months=months,
                           counts=counts)

@app.route("/track", methods=["GET", "POST"])
def track_feedback():
    feedback = None
    if request.method == "POST":
        ref = request.form["reference"]
        contact = request.form.get("contact")
        query = Feedback.query.filter_by(reference=ref)
        if contact:
            query = query.filter((Feedback.contact_email == contact) | (Feedback.contact_phone == contact))
        feedback = query.first()
        if not feedback:
            flash("No feedback found with that reference and contact info.")
    return render_template("track.html", feedback=feedback)

# -----------------------------------------------------------------
# PROTECTED ROUTES (Staff)
# -----------------------------------------------------------------
@app.route("/action_times", methods=["GET", "POST"])
def action_times():
    if request.method == "POST":
        fb_id = request.form.get("feedback_id")
        action_time_str = request.form.get("action_taken_at")
        if fb_id and action_time_str:
            fb = Feedback.query.get(fb_id)
            if fb:
                try:
                    fb.action_taken_at = datetime.strptime(action_time_str, "%Y-%m-%dT%H:%M")
                    db.session.commit()
                    flash("Action time updated successfully.", "success")
                except ValueError:
                    flash("Invalid datetime format.", "error")
        return redirect(url_for("action_times"))

    complaints = Feedback.query.filter_by(type='complaint').order_by(Feedback.created_at.desc()).all()
    for c in complaints:
        if c.action_taken_at:
            delta = c.action_taken_at - c.created_at
            c.action_hours = round(delta.total_seconds() / 3600, 1)
            c.within_48 = c.action_hours <= 48
        else:
            c.action_hours = None
            c.within_48 = None
    return render_template("action_times.html", complaints=complaints)

@app.route("/staff")
def staff_dashboard():
    return render_template("dashboard.html")

@app.route("/api/dashboard_stats")
def dashboard_stats():
    total = Feedback.query.count()
    unresolved = Feedback.query.filter(Feedback.final_status.in_(['Pending', None])).count()
    resolved = Feedback.query.filter_by(final_status='Resolved').count()
    implemented = Feedback.query.filter_by(final_status='Implemented').count()
    referred = Feedback.query.filter_by(final_status='Referred').count()
    recent = Feedback.query.order_by(Feedback.created_at.desc()).limit(5).all()
    recent_list = [{
        'id': f.id,
        'reference': f.reference,
        'issue': f.issue_received,  # full text
        'type': f.type,
        'date': f.created_at.strftime('%Y-%m-%d')
    } for f in recent]
    return jsonify({
        'total': total,
        'unresolved': unresolved,
        'resolved': resolved,
        'implemented': implemented,
        'referred': referred,
        'recent': recent_list
    })

@app.route("/submit", methods=["GET", "POST"])
def submit():
    if request.method == "POST":
        uploaded_files = request.files.getlist("attachments")
        filenames = []
        for file in uploaded_files:
            if file and file.filename:
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(filepath)
                filenames.append(filename)

        fb = Feedback(
            issue_received=request.form["issue"],
            type=request.form["type"],
            mechanism=request.form.get("mechanism"),
            recommendation=request.form.get("recommendation"),
            contact_email=request.form.get("contact_email"),
            contact_phone=request.form.get("contact_phone")
        )
        db.session.add(fb)
        db.session.flush()

        category_ids = request.form.getlist("categories")
        if category_ids:
            fb.categories = Category.query.filter(Category.id.in_(category_ids)).all()

        for fname in filenames:
            att = Attachment(feedback_id=fb.id, filename=fname, filepath=fname)
            db.session.add(att)

        db.session.commit()
        log_audit('feedback_submitted', f'Feedback {fb.reference} submitted by staff')
        flash("Feedback submitted successfully!", "success")
        return redirect(url_for("submit"))
    categories = Category.query.all()
    return render_template("submit.html", categories=categories)

@app.route("/summary_log")
def summary_log():
    query = Feedback.query
    type_filter = request.args.get('type')
    if type_filter and type_filter != 'all':
        query = query.filter_by(type=type_filter)
    status_filter = request.args.get('status')
    if status_filter and status_filter != 'all':
        query = query.filter_by(final_status=status_filter)
    category_id = request.args.get('category', type=int)
    if category_id:
        query = query.join(Feedback.categories).filter(Category.id == category_id)
    keyword = request.args.get('q')
    if keyword:
        query = query.filter(Feedback.issue_received.contains(keyword) |
                             Feedback.recommendation.contains(keyword))
    start_date = request.args.get('start_date')
    if start_date:
        start = datetime.strptime(start_date, '%Y-%m-%d')
        query = query.filter(Feedback.created_at >= start)
    end_date = request.args.get('end_date')
    if end_date:
        end = datetime.strptime(end_date, '%Y-%m-%d')
        query = query.filter(Feedback.created_at <= end)
    feedbacks = query.order_by(Feedback.created_at.desc()).all()
    categories = Category.query.all()
    return render_template("summary_log.html", feedbacks=feedbacks, categories=categories)

@app.route("/update_feedback/<int:fb_id>", methods=["POST"])
def update_feedback(fb_id):
    fb = Feedback.query.get_or_404(fb_id)
    old_status = fb.final_status
    first_action = request.form.get("first_action")
    if first_action and not fb.first_action:
        fb.first_action = first_action
        fb.action_taken_at = datetime.now()
    fb.final_comment = request.form.get("final_comment")
    new_status = request.form.get("final_status")
    if new_status and new_status != old_status:
        fb.final_status = new_status
        fb.action_timestamp = datetime.now()
        if new_status == "Resolved":
            fb.resolved_at = datetime.now()
    category_ids = request.form.getlist("categories")
    if category_ids:
        fb.categories = Category.query.filter(Category.id.in_(category_ids)).all()
    db.session.commit()
    history = FeedbackHistory(
        feedback_id=fb.id,
        user_id=session.get('user_id'),
        action='status_updated' if new_status else 'updated',
        old_status=old_status,
        new_status=new_status,
        comment=request.form.get("history_comment")
    )
    db.session.add(history)
    db.session.commit()
    log_audit('feedback_updated', f'Feedback {fb.reference} updated')
    return redirect(url_for("summary_log"))

@app.route("/categories", methods=["GET", "POST"])
def manage_categories():
    if request.method == "POST":
        name = request.form["name"]
        if name:
            cat = Category(name=name)
            db.session.add(cat)
            db.session.commit()
            log_audit('category_created', f'Category {name}')
            return redirect(url_for("manage_categories"))
    categories = Category.query.all()
    return render_template("categories.html", categories=categories)

@app.route("/category/delete/<int:cat_id>")
def delete_category(cat_id):
    cat = Category.query.get_or_404(cat_id)
    db.session.delete(cat)
    db.session.commit()
    log_audit('category_deleted', f'Category {cat.name}')
    return redirect(url_for("manage_categories"))

@app.route("/feedback/<int:fb_id>/history")
def feedback_history(fb_id):
    fb = Feedback.query.get_or_404(fb_id)
    return render_template("feedback_history.html", feedback=fb)

@app.route("/export_csv")
def export_csv():
    feedbacks = Feedback.query.all()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['ID', 'Reference', 'Type', 'Issue', 'Created', 'Final Status', 'Contact Email', 'Contact Phone', 'Categories'])
    for fb in feedbacks:
        cats = ', '.join([c.name for c in fb.categories])
        writer.writerow([fb.id, fb.reference, fb.type, fb.issue_received, fb.created_at, fb.final_status, fb.contact_email, fb.contact_phone, cats])
    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8')),
        mimetype='text/csv',
        as_attachment=True,
        download_name=f'feedback_export_{datetime.now().strftime("%Y%m%d")}.csv'
    )

@app.route("/export_docx")
def export_docx():
    feedbacks = Feedback.query.all()
    doc = Document()
    doc.add_heading('Feedback Export', 0)
    for fb in feedbacks:
        doc.add_heading(f'Reference: {fb.reference}', level=1)
        doc.add_paragraph(f'Type: {fb.type}')
        doc.add_paragraph(f'Issue: {fb.issue_received}')
        doc.add_paragraph(f'Recommendation: {fb.recommendation}')
        doc.add_paragraph(f'Status: {fb.final_status}')
        doc.add_paragraph(f'Created: {fb.created_at}')
        doc.add_paragraph('---')
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return send_file(
        file_stream,
        as_attachment=True,
        download_name=f'feedback_export_{datetime.now().strftime("%Y%m%d")}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.route("/analysis")
def analysis():
    return render_template("analysis.html", now=datetime.now())

@app.route("/deep_analysis_data")
def deep_analysis_data():
    month = request.args.get("month", type=int)
    quarter = request.args.get("quarter")
    year = request.args.get("year", type=int)
    q = Feedback.query
    if year:
        q = q.filter(db.extract("year", Feedback.created_at) == year)
    if month:
        q = q.filter(db.extract("month", Feedback.created_at) == month)
    if quarter:
        if quarter == "Q1":
            q = q.filter(db.extract("month", Feedback.created_at).between(1,3))
        elif quarter == "Q2":
            q = q.filter(db.extract("month", Feedback.created_at).between(4,6))
        elif quarter == "Q3":
            q = q.filter(db.extract("month", Feedback.created_at).between(7,9))
        elif quarter == "Q4":
            q = q.filter(db.extract("month", Feedback.created_at).between(10,12))
    data = q.all()
    total_feedback = len(data)
    total_complaints = sum(1 for f in data if f.type=="complaint")
    total_suggestions = sum(1 for f in data if f.type=="suggestion")
    total_compliments = sum(1 for f in data if f.type=="compliment")
    urgent_suggestions = [f for f in data if f.type=="suggestion" and "urgent" in f.issue_received.lower()]
    nonurgent_suggestions = [f for f in data if f.type=="suggestion" and "urgent" not in f.issue_received.lower()]
    urgent_complaints = [f for f in data if f.type=="complaint" and "urgent" in f.issue_received.lower()]
    nonurgent_complaints = [f for f in data if f.type=="complaint" and "urgent" not in f.issue_received.lower()]
    urgent_suggestion_rate = (sum(1 for f in urgent_suggestions if f.final_status=="Implemented") / len(urgent_suggestions)*100) if urgent_suggestions else 0
    urgent_complaint_rate = (sum(1 for f in urgent_complaints if f.final_status=="Resolved") / len(urgent_complaints)*100) if urgent_complaints else 0
    compliments_maintained = sum(1 for f in data if f.type=="compliment" and f.final_status=="Maintained")
    compliments_deviated = sum(1 for f in data if f.type=="compliment" and f.final_status=="Deviated")
    compliment_total = compliments_maintained + compliments_deviated
    compliment_maintain_rate = (compliments_maintained/compliment_total*100) if compliment_total else 0
    compliment_deviation_rate = (compliments_deviated/compliment_total*100) if compliment_total else 0
    referrals = sum(1 for f in data if f.final_status=="Referred")
    resolution_buckets = { "Same Day":0, "1-3 Days":0, "4-7 Days":0, "7+ Days":0 }
    for f in data:
        if f.type == "complaint" and f.resolved_at:
            days = (f.resolved_at - f.created_at).days
            if days==0:
                resolution_buckets["Same Day"] +=1
            elif 1<=days<=3:
                resolution_buckets["1-3 Days"] +=1
            elif 4<=days<=7:
                resolution_buckets["4-7 Days"] +=1
            else:
                resolution_buckets["7+ Days"] +=1
    mechanism_counts = {}
    for f in data:
        mechanism_counts[f.mechanism] = mechanism_counts.get(f.mechanism,0) + 1

    action_within_48h = 0
    action_over_48h = 0
    for f in data:
        if f.type == "complaint" and f.action_taken_at:
            delta = f.action_taken_at - f.created_at
            hours = delta.total_seconds() / 3600
            if hours <= 48:
                action_within_48h += 1
            else:
                action_over_48h += 1

    carried_over = 0
    if year and month:
        filter_start = datetime(year, month, 1)
        carried_over = Feedback.query.filter(
            Feedback.type == "complaint",
            Feedback.final_status.in_([None, "Pending"]),
            Feedback.created_at < filter_start
        ).count()
    elif year and quarter:
        if quarter == "Q1":
            filter_start = datetime(year, 1, 1)
        elif quarter == "Q2":
            filter_start = datetime(year, 4, 1)
        elif quarter == "Q3":
            filter_start = datetime(year, 7, 1)
        else:
            filter_start = datetime(year, 10, 1)
        carried_over = Feedback.query.filter(
            Feedback.type == "complaint",
            Feedback.final_status.in_([None, "Pending"]),
            Feedback.created_at < filter_start
        ).count()
    elif year:
        filter_start = datetime(year, 1, 1)
        carried_over = Feedback.query.filter(
            Feedback.type == "complaint",
            Feedback.final_status.in_([None, "Pending"]),
            Feedback.created_at < filter_start
        ).count()
    else:
        carried_over = sum(1 for f in data if f.type == "complaint" and f.final_status in (None, "Pending"))

    category_counts = {}
    for fb in data:
        for cat in fb.categories:
            category_counts[cat.name] = category_counts.get(cat.name, 0) + 1

    result = {
        "total_feedback": total_feedback,
        "total_complaints": total_complaints,
        "total_suggestions": total_suggestions,
        "total_compliments": total_compliments,
        "urgent_suggestions": len(urgent_suggestions),
        "nonurgent_suggestions": len(nonurgent_suggestions),
        "urgent_complaints": len(urgent_complaints),
        "nonurgent_complaints": len(nonurgent_complaints),
        "urgent_suggestion_rate": urgent_suggestion_rate,
        "urgent_complaint_rate": urgent_complaint_rate,
        "compliments_maintained": compliments_maintained,
        "compliments_deviated": compliments_deviated,
        "compliment_maintain_rate": compliment_maintain_rate,
        "compliment_deviation_rate": compliment_deviation_rate,
        "total_referrals": referrals,
        "resolution_buckets": resolution_buckets,
        "mechanism_counts": mechanism_counts,
        "action_within_48h": action_within_48h,
        "action_over_48h": action_over_48h,
        "carried_over": carried_over,
        "category_counts": category_counts
    }
    return jsonify(result)

@app.route("/deep_analysis")
def deep_analysis():
    return render_template("deep_analysis.html", now=datetime.now())

@app.route("/deep_trends")
def deep_trends():
    year = request.args.get("year", default=datetime.now().year, type=int)
    feedback_total_by_month = {m: 0 for m in range(1,13)}
    complaint_by_month = {m: {"resolved": 0, "total": 0} for m in range(1,13)}
    suggestion_by_month = {m: {"implemented": 0, "total": 0} for m in range(1,13)}
    compliment_by_month = {m: {"maintained": 0, "deviated": 0} for m in range(1,13)}
    urgent_complaint_by_month = {m: {"resolved": 0, "urgent": 0} for m in range(1,13)}
    urgent_suggestion_by_month = {m: {"implemented": 0, "urgent": 0} for m in range(1,13)}
    mechanism_by_month = {m: {} for m in range(1,13)}
    referral_by_month = {m: 0 for m in range(1,13)}
    action_within_48h_by_month = {m: 0 for m in range(1,13)}
    action_over_48h_by_month = {m: 0 for m in range(1,13)}
    carried_over_by_month = {m: 0 for m in range(1,13)}

    data = Feedback.query.filter(db.extract("year", Feedback.created_at) == year).all()
    for fb in data:
        m = fb.created_at.month
        feedback_total_by_month[m] += 1
        if fb.type == "complaint":
            complaint_by_month[m]["total"] += 1
            if fb.final_status == "Resolved":
                complaint_by_month[m]["resolved"] += 1
            if "urgent" in fb.issue_received.lower():
                urgent_complaint_by_month[m]["urgent"] += 1
                if fb.final_status == "Resolved":
                    urgent_complaint_by_month[m]["resolved"] += 1
            if fb.final_status == "Referred":
                referral_by_month[m] += 1
            if fb.action_taken_at:
                delta = fb.action_taken_at - fb.created_at
                hours = delta.total_seconds() / 3600
                if hours <= 48:
                    action_within_48h_by_month[m] += 1
                else:
                    action_over_48h_by_month[m] += 1
        elif fb.type == "suggestion":
            suggestion_by_month[m]["total"] += 1
            if fb.final_status == "Implemented":
                suggestion_by_month[m]["implemented"] += 1
            if "urgent" in fb.issue_received.lower():
                urgent_suggestion_by_month[m]["urgent"] += 1
                if fb.final_status == "Implemented":
                    urgent_suggestion_by_month[m]["implemented"] += 1
        elif fb.type == "compliment":
            if fb.final_status == "Maintained":
                compliment_by_month[m]["maintained"] += 1
            elif fb.final_status == "Deviated":
                compliment_by_month[m]["deviated"] += 1
        mech = fb.mechanism
        if mech:
            mechanism_by_month[m][mech] = mechanism_by_month[m].get(mech, 0) + 1

    for m in range(1,13):
        month_start = datetime(year, m, 1)
        carried_over_by_month[m] = Feedback.query.filter(
            Feedback.type == "complaint",
            Feedback.final_status.in_([None, "Pending"]),
            Feedback.created_at < month_start
        ).count()

    return render_template(
        "deep_trends.html",
        now=datetime.now(),
        year=year,
        feedback_total_by_month=feedback_total_by_month,
        complaint_by_month=complaint_by_month,
        suggestion_by_month=suggestion_by_month,
        compliment_by_month=compliment_by_month,
        urgent_complaint_by_month=urgent_complaint_by_month,
        urgent_suggestion_by_month=urgent_suggestion_by_month,
        mechanism_by_month=mechanism_by_month,
        referral_by_month=referral_by_month,
        action_within_48h_by_month=action_within_48h_by_month,
        action_over_48h_by_month=action_over_48h_by_month,
        carried_over_by_month=carried_over_by_month,
        max=max
    )

@app.route("/current_month")
def current_month():
    now = datetime.now()
    year = now.year
    month = now.month
    feedbacks = Feedback.query.filter(
        db.extract('year', Feedback.created_at) == year,
        db.extract('month', Feedback.created_at) == month
    ).all()

    total = len(feedbacks)
    complaints = [f for f in feedbacks if f.type == 'complaint']
    suggestions = [f for f in feedbacks if f.type == 'suggestion']
    compliments = [f for f in feedbacks if f.type == 'compliment']

    urgent_suggestions = [f for f in suggestions if 'urgent' in f.issue_received.lower()]
    urgent_complaints = [f for f in complaints if 'urgent' in f.issue_received.lower()]

    action_within_48h = 0
    action_over_48h = 0
    for c in complaints:
        if c.action_taken_at:
            delta = c.action_taken_at - c.created_at
            hours = delta.total_seconds() / 3600
            if hours <= 48:
                action_within_48h += 1
            else:
                action_over_48h += 1

    resolved_complaints = [c for c in complaints if c.final_status == 'Resolved']
    referred_complaints = [c for c in complaints if c.final_status == 'Referred']
    pending_complaints = [c for c in complaints if c.final_status in (None, 'Pending')]

    carried_over = Feedback.query.filter(
        Feedback.type == 'complaint',
        Feedback.final_status.in_([None, 'Pending']),
        Feedback.created_at < datetime(year, month, 1)
    ).count()

    category_counts = {}
    for fb in feedbacks:
        for cat in fb.categories:
            category_counts[cat.name] = category_counts.get(cat.name, 0) + 1

    resolution_buckets = {"Same Day":0, "1-3 Days":0, "4-7 Days":0, "7+ Days":0}
    for c in complaints:
        if c.resolved_at:
            days = (c.resolved_at - c.created_at).days
            if days == 0:
                resolution_buckets["Same Day"] += 1
            elif 1 <= days <= 3:
                resolution_buckets["1-3 Days"] += 1
            elif 4 <= days <= 7:
                resolution_buckets["4-7 Days"] += 1
            else:
                resolution_buckets["7+ Days"] += 1

    return render_template("current_month.html",
                           now=now,
                           total=total,
                           complaints=len(complaints),
                           suggestions=len(suggestions),
                           compliments=len(compliments),
                           urgent_suggestions=len(urgent_suggestions),
                           urgent_complaints=len(urgent_complaints),
                           action_within_48h=action_within_48h,
                           action_over_48h=action_over_48h,
                           resolved=len(resolved_complaints),
                           referred=len(referred_complaints),
                           pending=len(pending_complaints),
                           carried_over=carried_over,
                           category_counts=category_counts,
                           resolution_buckets=resolution_buckets)

@app.route("/about")
def about():
    return render_template("about.html")

@app.route("/contact")
def contact():
    return render_template("contact.html")

@app.route("/help")
def help():
    return render_template("help.html")

@app.route("/all_feedback")
def all_feedback():
    data = Feedback.query.all()
    return jsonify([{
        "id": fb.id,
        "reference": fb.reference,
        "issue_received": fb.issue_received,
        "created_at": fb.created_at.isoformat(),
        "type": fb.type,
        "mechanism": fb.mechanism,
        "recommendation": fb.recommendation,
        "final_status": fb.final_status,
        "resolved_at": fb.resolved_at.isoformat() if fb.resolved_at else None,
        "contact_email": fb.contact_email,
        "contact_phone": fb.contact_phone,
        "categories": [c.name for c in fb.categories]
    } for fb in data])

# -----------------------------------------------------------------
# ADMIN-ONLY ROUTES
# -----------------------------------------------------------------
@app.route('/admin/users', methods=['GET', 'POST'])
@admin_required
def admin_users():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        role = request.form.get('role', 'viewer')
        if not username or not password:
            flash('Username and password required.', 'error')
            return redirect(url_for('admin_users'))
        existing = User.query.filter_by(username=username).first()
        if existing:
            flash('Username already exists.', 'error')
            return redirect(url_for('admin_users'))
        user = User(username=username, role=role)
        user.set_password(password)
        db.session.add(user)
        db.session.commit()
        log_audit('user_created', f'User {username} created by admin')
        flash(f'User {username} created successfully.', 'success')
        return redirect(url_for('admin_users'))
    users = User.query.all()
    return render_template('admin_users.html', users=users)

@app.route('/admin/user/delete/<int:user_id>')
@admin_required
def admin_user_delete(user_id):
    user = User.query.get_or_404(user_id)
    if user.id == session.get('user_id'):
        flash('You cannot delete yourself.', 'error')
        return redirect(url_for('admin_users'))
    db.session.delete(user)
    db.session.commit()
    log_audit('user_deleted', f'User {user.username} deleted by admin')
    flash(f'User {user.username} deleted.', 'success')
    return redirect(url_for('admin_users'))

@app.route('/admin/feedback/delete/<int:fb_id>', methods=['POST'])
@admin_required
def admin_feedback_delete(fb_id):
    fb = Feedback.query.get_or_404(fb_id)
    log_audit('feedback_deleted', f'Feedback {fb.reference} deleted by admin')
    db.session.delete(fb)
    db.session.commit()
    flash(f'Feedback {fb.reference} deleted.', 'success')
    return redirect(request.referrer or url_for('summary_log'))

@app.route('/admin/backup')
@admin_required
def admin_backup():
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Users
        users = User.query.all()
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(['id', 'username', 'role', 'password_hash'])
        for u in users:
            writer.writerow([u.id, u.username, u.role, u.password_hash])
        zip_file.writestr('users.csv', csv_buffer.getvalue())

        # Categories
        categories = Category.query.all()
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(['id', 'name'])
        for c in categories:
            writer.writerow([c.id, c.name])
        zip_file.writestr('categories.csv', csv_buffer.getvalue())

        # Feedback
        feedbacks = Feedback.query.all()
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow([
            'id', 'reference', 'issue_received', 'created_at', 'type', 'mechanism',
            'recommendation', 'first_action', 'action_taken_at', 'implementation_status',
            'final_comment', 'final_status', 'action_timestamp', 'resolved_at',
            'contact_email', 'contact_phone', 'categories'
        ])
        for fb in feedbacks:
            cats = ', '.join([c.name for c in fb.categories])
            writer.writerow([
                fb.id, fb.reference, fb.issue_received, fb.created_at, fb.type, fb.mechanism,
                fb.recommendation, fb.first_action, fb.action_taken_at, fb.implementation_status,
                fb.final_comment, fb.final_status, fb.action_timestamp, fb.resolved_at,
                fb.contact_email, fb.contact_phone, cats
            ])
        zip_file.writestr('feedback.csv', csv_buffer.getvalue())

        # FeedbackHistory
        histories = FeedbackHistory.query.all()
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(['id', 'feedback_id', 'user_id', 'action', 'old_status', 'new_status', 'comment', 'timestamp'])
        for h in histories:
            writer.writerow([h.id, h.feedback_id, h.user_id, h.action, h.old_status, h.new_status, h.comment, h.timestamp])
        zip_file.writestr('feedback_history.csv', csv_buffer.getvalue())

        # Attachments
        attachments = Attachment.query.all()
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(['id', 'feedback_id', 'filename', 'filepath', 'uploaded_at'])
        for a in attachments:
            writer.writerow([a.id, a.feedback_id, a.filename, a.filepath, a.uploaded_at])
        zip_file.writestr('attachments.csv', csv_buffer.getvalue())

        # AuditLog
        audit_logs = AuditLog.query.all()
        csv_buffer = io.StringIO()
        writer = csv.writer(csv_buffer)
        writer.writerow(['id', 'user_id', 'action', 'details', 'timestamp'])
        for al in audit_logs:
            writer.writerow([al.id, al.user_id, al.action, al.details, al.timestamp])
        zip_file.writestr('audit_log.csv', csv_buffer.getvalue())

        # Uploaded files
        upload_dir = app.config['UPLOAD_FOLDER']
        if os.path.exists(upload_dir):
            for root, dirs, files in os.walk(upload_dir):
                for file in files:
                    file_path = os.path.join(root, file)
                    arcname = os.path.join('uploads', file)
                    zip_file.write(file_path, arcname)

    zip_buffer.seek(0)
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return send_file(
        zip_buffer,
        mimetype='application/zip',
        as_attachment=True,
        download_name=f'backup_{timestamp}.zip'
    )

@app.route('/admin/restore', methods=['GET', 'POST'])
@admin_required
def admin_restore():
    if request.method == 'POST':
        if 'backup_file' not in request.files:
            flash('No file selected.', 'error')
            return redirect(url_for('admin_restore'))
        file = request.files['backup_file']
        if file.filename == '':
            flash('No file selected.', 'error')
            return redirect(url_for('admin_restore'))
        if not file.filename.endswith('.zip'):
            flash('Only ZIP files allowed.', 'error')
            return redirect(url_for('admin_restore'))

        temp_dir = tempfile.mkdtemp()
        try:
            zip_path = os.path.join(temp_dir, 'backup.zip')
            file.save(zip_path)
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # Clear existing data in correct order
            db.session.execute(text('DELETE FROM feedback_categories'))
            db.session.query(Attachment).delete()
            db.session.query(FeedbackHistory).delete()
            db.session.query(Feedback).delete()
            db.session.query(Category).delete()
            db.session.query(User).delete()
            db.session.query(AuditLog).delete()
            db.session.commit()

            # Import categories
            categories_csv = os.path.join(temp_dir, 'categories.csv')
            if os.path.exists(categories_csv):
                with open(categories_csv, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    next(reader)
                    for row in reader:
                        if len(row) >= 2:
                            cat = Category(name=row[1])
                            db.session.add(cat)
                    db.session.commit()

            # Import users
            users_csv = os.path.join(temp_dir, 'users.csv')
            if os.path.exists(users_csv):
                with open(users_csv, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    next(reader)
                    for row in reader:
                        if len(row) >= 4:
                            user = User(
                                id=int(row[0]),
                                username=row[1],
                                role=row[2],
                                password_hash=row[3]
                            )
                            db.session.add(user)
                    db.session.commit()

            # Import feedback
            feedbacks_csv = os.path.join(temp_dir, 'feedback.csv')
            if os.path.exists(feedbacks_csv):
                with open(feedbacks_csv, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    next(reader)
                    for row in reader:
                        if len(row) >= 17:
                            fb = Feedback(
                                id=int(row[0]),
                                reference=row[1],
                                issue_received=row[2],
                                created_at=parse_datetime(row[3]),
                                type=row[4],
                                mechanism=row[5] if row[5] else None,
                                recommendation=row[6] if row[6] else None,
                                first_action=row[7] if row[7] else None,
                                action_taken_at=parse_datetime(row[8]),
                                implementation_status=row[9] if row[9] else None,
                                final_comment=row[10] if row[10] else None,
                                final_status=row[11] if row[11] else None,
                                action_timestamp=parse_datetime(row[12]),
                                resolved_at=parse_datetime(row[13]),
                                contact_email=row[14] if row[14] else None,
                                contact_phone=row[15] if row[15] else None
                            )
                            db.session.add(fb)
                    db.session.commit()

            # Assign categories (many-to-many)
            with open(feedbacks_csv, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                next(reader)
                for row in reader:
                    if len(row) >= 17:
                        fb_id = int(row[0])
                        cats_str = row[16]
                        if cats_str:
                            fb = Feedback.query.get(fb_id)
                            if fb:
                                cat_names = [c.strip() for c in cats_str.split(',')]
                                cats = Category.query.filter(Category.name.in_(cat_names)).all()
                                fb.categories = cats
                db.session.commit()

            # Import feedback history
            history_csv = os.path.join(temp_dir, 'feedback_history.csv')
            if os.path.exists(history_csv):
                with open(history_csv, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    next(reader)
                    for row in reader:
                        if len(row) >= 8:
                            h = FeedbackHistory(
                                id=int(row[0]),
                                feedback_id=int(row[1]),
                                user_id=int(row[2]) if row[2] else None,
                                action=row[3],
                                old_status=row[4] if row[4] else None,
                                new_status=row[5] if row[5] else None,
                                comment=row[6] if row[6] else None,
                                timestamp=parse_datetime(row[7])
                            )
                            db.session.add(h)
                    db.session.commit()

            # Import attachments
            attachments_csv = os.path.join(temp_dir, 'attachments.csv')
            if os.path.exists(attachments_csv):
                with open(attachments_csv, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    next(reader)
                    for row in reader:
                        if len(row) >= 5:
                            a = Attachment(
                                id=int(row[0]),
                                feedback_id=int(row[1]),
                                filename=row[2],
                                filepath=row[3],
                                uploaded_at=parse_datetime(row[4])
                            )
                            db.session.add(a)
                    db.session.commit()

            # Import audit log
            audit_csv = os.path.join(temp_dir, 'audit_log.csv')
            if os.path.exists(audit_csv):
                with open(audit_csv, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    next(reader)
                    for row in reader:
                        if len(row) >= 5:
                            al = AuditLog(
                                id=int(row[0]),
                                user_id=int(row[1]) if row[1] else None,
                                action=row[2],
                                details=row[3] if row[3] else None,
                                timestamp=parse_datetime(row[4])
                            )
                            db.session.add(al)
                    db.session.commit()

            # Restore uploaded files
            uploads_source = os.path.join(temp_dir, 'uploads')
            if os.path.exists(uploads_source):
                upload_dest = app.config['UPLOAD_FOLDER']
                if os.path.exists(upload_dest):
                    shutil.rmtree(upload_dest)
                shutil.copytree(uploads_source, upload_dest)

            log_audit('system_restore', 'System restored from backup')
            session.clear()
            flash('System restored successfully. Please log in again with restored credentials.', 'success')
            return redirect(url_for('login'))

        except Exception as e:
            db.session.rollback()
            flash(f'Restore failed: {str(e)}', 'error')
            return redirect(url_for('admin_restore'))
        finally:
            shutil.rmtree(temp_dir, ignore_errors=True)

    return render_template('admin_restore.html')

# -----------------------------------------------------------------
# AUDIT LOG VIEW
# -----------------------------------------------------------------
@app.route('/admin/audit_log')
@admin_required
def admin_audit_log():
    page = request.args.get('page', 1, type=int)
    per_page = 50
    logs = AuditLog.query.options(joinedload(AuditLog.user)).order_by(AuditLog.timestamp.desc()).paginate(page=page, per_page=per_page)
    return render_template('admin_audit_log.html', logs=logs)

# -----------------------------------------------------------------
# INITIALIZATION (development only)
# -----------------------------------------------------------------
if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", use_reloader=False)